# 导入模块
import asyncio
import json
import os
from datetime import datetime, timedelta

import yaml
from docx2pdf import convert
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ROW_HEIGHT_RULE
from ai.backend.util.db.configuration.path import get_config_path
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.util.db_tool_weekly import DbToolsCsv
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.util.util import csv_to_json
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.util.agent import ask_question
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.output.output_path import get_output_path


def load_custody(db, brand, country=None):
    # 从 YAML 文件加载数据库信息
    Brand_path = os.path.join(get_config_path(), 'Brand.yml')
    with open(Brand_path, 'r') as file:
        Brand_data = yaml.safe_load(file)
    brand_info = Brand_data.get(db, {})
    brand_info = brand_info.get(brand, {})
    if country:
        country_info = brand_info.get(country, {})
        custody = country_info.get('custody', brand_info.get('default', {}).get('custody'))
        if custody == 'T':
            return country_info.get('parent_asins', brand_info.get('default', {}).get('parent_asins'))
        return False  # custody 不是 'T' 时返回空字典
    return False


def set_font_size(cell, pt_size):
    """设置单元格字体大小"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(pt_size)


def get_font_size1(value):
    """根据float值的长度确定字体大小"""
    if isinstance(value, float):
        length = len(str(int(value)))
        if length >= 10:
            return 6
        elif length == 9:
            return 7
        elif length == 8:
            return 8
        elif length <= 7:
            return 9
        else:
            return 6
    return 6


def get_font_size2(value):
    """根据float值的长度确定字体大小"""
    if isinstance(value, float):
        length = len(str(int(value)))
        if length >= 9:
            return 6
        elif length == 8:
            return 7
        elif length == 7:
            return 7
        elif length <= 6:
            return 8
        else:
            return 6
    return 6


def determine_max_font_size(csv_path,code):
    """读取CSV文件并返回根据float数据长度确定的最大字体大小"""
    df = pd.read_csv(csv_path, encoding='utf-8-sig')

    max_font_size = 9  # 默认字体大小

    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            value = df.iloc[i, j]
            if isinstance(value, (float, np.float64)) and not pd.isna(value):
                if code == 1:
                    font_size = get_font_size1(value)
                elif code == 2:
                    font_size = get_font_size2(value)
                if font_size < max_font_size:
                    max_font_size = font_size
    print(max_font_size)
    return max_font_size


def insert_csv_to_docx(doc, csv_path, pt):
    """将CSV文件插入到给定的docx文档中"""
    df = pd.read_csv(csv_path, encoding='utf-8-sig')
    row = df.shape[0] + 1  # 行数，加标题栏
    col = df.shape[1]  # 列数

    # 设置表格样式
    table = doc.add_table(rows=row, cols=col, style="Table Grid")

    # 写入列标签
    for i in range(col):
        cell = table.cell(0, i)
        cell.text = str(df.columns[i])
        set_font_size(cell, pt)

        # 写入数据
    for i in range(1, row):
        for j in range(col):
            cell = table.cell(i, j)
            value = df.iloc[i - 1, j]
            if pd.isna(value):  # 检查是否为 NaN
                cell.text = ' '
            elif isinstance(value, str):
                cell.text = value
            elif isinstance(value, (int, np.integer)):
                cell.text = str(int(value))  # 确保整数类型为 int
            elif isinstance(value, float):
                cell.text = '{:.0f}'.format(value)
            set_font_size(cell, pt)

    #table.autofit = True  # 表格自动适应窗口大小
    # table.rows.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    # table.style.font.name = u'楷体'  # 设置字体格式
    # table.style.font.size = Pt(pt)  # 设置字体大小
    return table


def set_table_width(table, width_in_inches):
    """设置整个表格的宽度（单位：英寸）。"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    # 如果 tblW 不存在，则创建一个新的
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)

    tblW.set(qn('w:type'), 'dxa')  # 设置宽度类型为 'dxa'（twips）
    tblW.set(qn('w:w'), str(int(width_in_inches * 1440)))  # 设置宽度，1 英寸 = 1440 twips

    # 确保表格布局是固定的
    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')  # 设置表格布局为固定宽度


def create_summery(date,code):
    summery = asyncio.get_event_loop().run_until_complete(ask_question(date, code))
    return summery
    # return 'test'


def generate_docx(db, brand, market,start_date,end_date,period):
    # 新建文档对象按模板新建 word 文档文件，具有模板文件的所有格式
    doc = Document()
    # end_date = (datetime.today() - timedelta(days=2)).strftime('%m.%d')
    # start_date = (datetime.today() - timedelta(days=31)).strftime('%m.%d')
    # 增加标题：add_heading(self, text="", level=1):
    date1 = datetime.strptime(end_date, '%Y-%m-%d')
    date1_date = date1.strftime('%m 月 %d 日')
    date2 = datetime.strptime(start_date, '%Y-%m-%d')
    date2_date = date2.strftime('%m 月 %d 日')
    seven_day_before_date1 = (date1 - timedelta(days=period)).strftime('%Y-%m-%d')
    seven_day_before_date2 = (date2 - timedelta(days=period)).strftime('%Y-%m-%d')
    seven_days_before_date1 = (date1 - timedelta(days=period)).strftime('%m 月 %d 日')
    seven_days_before_date2 = (date2 - timedelta(days=period)).strftime('%m 月 %d 日')
    difference = (date1 - date2).days+1
    date = datetime.strptime(end_date, '%Y-%m-%d').date()
    year_start = datetime(date.year, 1, 1).date()
    days_passed = (date - year_start).days
    week_number = days_passed // 31 + 1
    doc.add_heading(f'客户名称：\n第{week_number}月月报（{date2_date}-{date1_date}）', 1)

    doc.add_heading('1. 目标与完成情况', 2)
    doc.add_heading('1.1 总体目标：', 3)
    table1 = doc.add_table(2, 7, style="Table Grid")
    table1.cell(0, 0).text = '总体目标'
    doc.add_heading(f'1.2 上月回顾（{seven_days_before_date2}-{seven_days_before_date1}）：', 3)
    csv_path1 = DbToolsCsv(db, brand, market).get_review_and_goals_data(market,seven_day_before_date2,seven_day_before_date1)
    table1 = insert_csv_to_docx(doc, csv_path1,8)
    # set_table_width(table1, 2)
    doc.add_heading(f'1.3 本月目标及实际（{date2_date}-{date1_date}）：', 3)
    csv_path2 = DbToolsCsv(db, brand, market).get_review_and_goals_data(market, start_date, end_date)
    table2 = insert_csv_to_docx(doc, csv_path2, 8)
    # 为表格对象增加行,add_row(self)
    table2.add_row()  # 只能逐行添加
    table2.cell(2, 0).text = '目标'
    doc.add_heading(f'1.4 完成情况分析：', 3)
    doc.add_paragraph().add_run('')

    doc.add_heading('2. 本月现状', 2)
    asin_info = load_custody(db, brand, market)
    if not asin_info:
        doc.add_heading(f'2.1 全店销售现状（{date2_date}-{date1_date}）：', 3)
        csv_path3 = DbToolsCsv(db, brand, market).get_store_sales_status(market, start_date, end_date, difference)
        insert_csv_to_docx(doc, csv_path3, 8)
        json1 = csv_to_json(csv_path3)
        translate_kw1 = create_summery(json1, 0)
        translate_kw2 = create_summery(json1, 1)
        doc.add_paragraph().add_run(f"1.{translate_kw1}")
        doc.add_paragraph().add_run(f"2.{translate_kw2}")
        doc.add_heading(f'2.2 托管Listing新旧计划对比（{date2_date}-{date1_date}）：', 3)
        csv_path4 = DbToolsCsv(db, brand, market).get_listing_comparison_data(market, start_date, end_date)
        insert_csv_to_docx(doc, csv_path4, 8)
        json2 = csv_to_json(csv_path4)
        translate_kw3 = create_summery(json2, 15)
        translate_kw4 = create_summery(json2, 16)
        doc.add_paragraph().add_run(f"1.{translate_kw3}")
        doc.add_paragraph().add_run(f"2.{translate_kw4}")
        doc.add_heading(f'2.3 托管listing明细（{date2_date}-{date1_date}）：', 3)
        if brand == 'LAPASA':
            csv_path5 = DbToolsCsv(db, brand, market).get_listing_ditial_data_lapasa(market, start_date, end_date)
        else:
            csv_path5 = DbToolsCsv(db, brand, market).get_listing_ditial_data(market, start_date, end_date)
        insert_csv_to_docx(doc, csv_path5, 8)
        json3 = csv_to_json(csv_path5)
        translate_kw5 = create_summery(json3, 18)
        translate_kw6 = create_summery(json3, 7)
        doc.add_paragraph().add_run(f"1.{translate_kw5}")
        doc.add_paragraph().add_run(f"2.{translate_kw6}")
        doc.add_heading(f'2.4 库存情况（{date2_date}-{date1_date}）：', 3)
        if brand == 'LAPASA':
            csv_path6 = DbToolsCsv(db, brand, market).get_inventory_data_lapasa(market, start_date, end_date)
            insert_csv_to_docx(doc, csv_path6, 8)
        else:
            csv_path6 = DbToolsCsv(db, brand, market).get_inventory_data(market, start_date, end_date)
            insert_csv_to_docx(doc, csv_path6, 8)
    else:
        doc.add_heading(f'2.1 托管Listing现状汇总（{date2_date}-{date1_date}）：', 3)
        csv_path3 = DbToolsCsv(db, brand, market).get_managed_listing_current_data(market, start_date, end_date, asin_info,seven_day_before_date2,seven_day_before_date1)
        insert_csv_to_docx(doc, csv_path3, 8)
        json1 = csv_to_json(csv_path3)
        translate_kw1 = create_summery(json1, 0)
        translate_kw2 = create_summery(json1, 17)
        doc.add_paragraph().add_run(f"1.{translate_kw1}")
        doc.add_paragraph().add_run(f"2.{translate_kw2}")
        doc.add_heading(f'2.2 托管Listing新旧计划对比（{date2_date}-{date1_date}）：', 3)
        csv_path4 = DbToolsCsv(db, brand, market).get_managed_listing_comparison_data(market, start_date, end_date, asin_info)
        insert_csv_to_docx(doc, csv_path4, 8)
        json2 = csv_to_json(csv_path4)
        translate_kw3 = create_summery(json2, 15)
        translate_kw4 = create_summery(json2, 16)
        doc.add_paragraph().add_run(f"1.{translate_kw3}")
        doc.add_paragraph().add_run(f"2.{translate_kw4}")
        doc.add_heading(f'2.3 托管listing明细（{date2_date}-{date1_date}）：', 3)
        csv_path5 = DbToolsCsv(db, brand, market).get_managed_listing_ditial_data(market, start_date, end_date, asin_info)
        insert_csv_to_docx(doc, csv_path5, 8)
        json3 = csv_to_json(csv_path5)
        translate_kw5 = create_summery(json3, 18)
        translate_kw6 = create_summery(json3, 7)
        doc.add_paragraph().add_run(f"1.{translate_kw5}")
        doc.add_paragraph().add_run(f"2.{translate_kw6}")
        doc.add_heading(f'2.4 库存情况（{date2_date}-{date1_date}）：', 3)
        csv_path6 = DbToolsCsv(db, brand, market).get_inventory_data_custody(market, start_date, end_date, asin_info)
        insert_csv_to_docx(doc, csv_path6, 8)


    doc.add_heading('3. 总结', 2)
    # result1 = translate_kw1 + "\n" + translate_kw2
    # translate_kw4 = create_summery(result1, 3)
    doc.add_heading('3.1 本月总结（整体总结）', 3)
    doc.add_heading('3.2 下月优化策略', 3)
    doc.add_heading('3.3 下月计划与目标', 3)
    doc.add_heading('3.3.1 下月目标设定', 4)
    doc.add_heading('3.3.2 行动计划', 4)
    doc.add_heading('4. 异常情况及客户协助', 2)


    # sections = doc.sections
    # print(sections)
    # default_section = sections[0]
    # print(default_section.top_margin.cm)  # 2.54
    # print(default_section.right_margin.cm)  # 3.175
    # print(default_section.bottom_margin.cm)  # 2.54
    # print(default_section.left_margin.cm)  # 3.175
    # default_section.top_margin = Cm(2.5)
    # default_section.right_margin = Cm(2)
    # default_section.bottom_margin = Cm(2.5)
    # default_section.left_margin = Cm(2)
    # # 增加分页符
    # doc.add_page_break()
    #
    # # 增加标题 API 分析， 只能设置 0-9 级标题
    # for i in range(0,10):
    #     doc.add_heading(f'标题{i}', i)

    docx_path = f'{brand}_{market}_2024年{week_number}月_({date2_date}-{date1_date}).docx'
    docx_path = os.path.join(get_output_path(), docx_path)
    # 保存文件
    doc.save(docx_path)
    return docx_path
    # convert(docx_path)

def generate_summary_docx(db, brand, market,start_date,end_date,period):
    # 新建文档对象按模板新建 word 文档文件，具有模板文件的所有格式
    doc = Document()
    # end_date = (datetime.today() - timedelta(days=2)).strftime('%m.%d')
    # start_date = (datetime.today() - timedelta(days=31)).strftime('%m.%d')
    # 增加标题：add_heading(self, text="", level=1):
    date1 = datetime.strptime(end_date, '%Y-%m-%d')
    date1_date = date1.strftime('%m 月 %d 日')
    date2 = datetime.strptime(start_date, '%Y-%m-%d')
    date2_date = date2.strftime('%m 月 %d 日')
    seven_day_before_date1 = (date1 - timedelta(days=period)).strftime('%Y-%m-%d')
    seven_day_before_date2 = (date2 - timedelta(days=period)).strftime('%Y-%m-%d')
    seven_days_before_date1 = (date1 - timedelta(days=period)).strftime('%m 月 %d 日')
    seven_days_before_date2 = (date2 - timedelta(days=period)).strftime('%m 月 %d 日')
    difference = (date1 - date2).days+1
    date = datetime.strptime(end_date, '%Y-%m-%d').date()
    year_start = datetime(date.year, 1, 1).date()
    days_passed = (date - year_start).days
    week_number = days_passed // 31 + 1
    doc.add_heading(f'客户名称：\n第{week_number}月月报（{date2_date}-{date1_date}）', 1)

    doc.add_heading('1. 目标与完成情况', 2)
    doc.add_heading('1.1 总体目标：', 3)
    table1 = doc.add_table(2, 7, style="Table Grid")
    table1.cell(0, 0).text = '总体目标'
    doc.add_heading(f'1.2 上月回顾（{seven_days_before_date2}-{seven_days_before_date1}）：', 3)
    csv_path1 = DbToolsCsv(db, brand, market).get_review_and_goals_data_summary(market,seven_day_before_date2,seven_day_before_date1)
    table1 = insert_csv_to_docx(doc, csv_path1,8)
    # set_table_width(table1, 2)
    doc.add_heading(f'1.3 本月目标及实际（{date2_date}-{date1_date}）：', 3)
    csv_path2 = DbToolsCsv(db, brand, market).get_review_and_goals_data_summary(market, start_date, end_date)
    table2 = insert_csv_to_docx(doc, csv_path2, 8)
    # 为表格对象增加行,add_row(self)
    table2.add_row()  # 只能逐行添加
    table2.cell(2, 0).text = '目标'
    doc.add_heading(f'1.4 完成情况分析：', 3)
    doc.add_paragraph().add_run('')

    doc.add_heading('2. 本月现状', 2)
    asin_info = load_custody(db, brand, market)
    if not asin_info:
        doc.add_heading(f'2.1 全店销售现状（{date2_date}-{date1_date}）：', 3)
        doc.add_heading(f'{market}数据表：', 4)
        csv_path3,countries = DbToolsCsv(db, brand, market).get_store_sales_status_summary(market, start_date, end_date, difference)
        insert_csv_to_docx(doc, csv_path3, 8)
        json1 = csv_to_json(csv_path3)
        translate_kw1 = create_summery(json1, 0)
        translate_kw2 = create_summery(json1, 1)
        doc.add_paragraph().add_run(f"1.{translate_kw1}")
        doc.add_paragraph().add_run(f"2.{translate_kw2}")
        for country in countries:
            doc.add_heading(f'{country}数据表：', 4)
            csv_path = DbToolsCsv(db, brand, country).get_store_sales_status(country, start_date, end_date, difference)
            insert_csv_to_docx(doc, csv_path, 8)
        doc.add_heading(f'2.2 托管Listing新旧计划对比（{date2_date}-{date1_date}）：', 3)
        csv_path4 = DbToolsCsv(db, brand, market).get_listing_comparison_data_summary(market, start_date, end_date)
        insert_csv_to_docx(doc, csv_path4, 8)
        json2 = csv_to_json(csv_path4)
        translate_kw3 = create_summery(json2, 15)
        translate_kw4 = create_summery(json2, 16)
        doc.add_paragraph().add_run(f"1.{translate_kw3}")
        doc.add_paragraph().add_run(f"2.{translate_kw4}")
        doc.add_heading(f'2.3 托管listing明细（{date2_date}-{date1_date}）：', 3)
        if brand == 'LAPASA':
            csv_path5 = DbToolsCsv(db, brand, market).get_listing_ditial_data_lapasa_summary(market, start_date, end_date)
        else:
            csv_path5 = DbToolsCsv(db, brand, market).get_listing_ditial_data_summary(market, start_date, end_date)
            print(csv_path5)
        print('test1')
        insert_csv_to_docx(doc, csv_path5, 8)
        print('test4')
        json3 = csv_to_json(csv_path5)
        translate_kw5 = create_summery(json3, 18)
        translate_kw6 = create_summery(json3, 7)
        doc.add_paragraph().add_run(f"1.{translate_kw5}")
        doc.add_paragraph().add_run(f"2.{translate_kw6}")
        doc.add_heading(f'2.4 库存情况（{date2_date}-{date1_date}）：', 3)
        if brand == 'LAPASA':
            print('test3')
            csv_path6 = DbToolsCsv(db, brand, market).get_inventory_data_lapasa(market, start_date, end_date)
            insert_csv_to_docx(doc, csv_path6, 8)
        else:
            print('test2')
            csv_path6 = DbToolsCsv(db, brand, market).get_inventory_data(market, start_date, end_date)
            insert_csv_to_docx(doc, csv_path6, 8)
    # else:
    #     doc.add_heading(f'2.1 托管Listing现状汇总（{date2_date}-{date1_date}）：', 3)
    #     csv_path3 = DbToolsCsv(brand, market).get_managed_listing_current_data(market, start_date, end_date, asin_info,seven_day_before_date2,seven_day_before_date1)
    #     insert_csv_to_docx(doc, csv_path3, 8)
    #     json1 = csv_to_json(csv_path3)
    #     translate_kw1 = create_summery(json1, 0)
    #     translate_kw2 = create_summery(json1, 17)
    #     doc.add_paragraph().add_run(f"1.{translate_kw1}")
    #     doc.add_paragraph().add_run(f"2.{translate_kw2}")
    #     doc.add_heading(f'2.2 托管Listing新旧计划对比（{date2_date}-{date1_date}）：', 3)
    #     csv_path4 = DbToolsCsv(brand, market).get_managed_listing_comparison_data(market, start_date, end_date, asin_info)
    #     insert_csv_to_docx(doc, csv_path4, 8)
    #     json2 = csv_to_json(csv_path4)
    #     translate_kw3 = create_summery(json2, 15)
    #     translate_kw4 = create_summery(json2, 16)
    #     doc.add_paragraph().add_run(f"1.{translate_kw3}")
    #     doc.add_paragraph().add_run(f"2.{translate_kw4}")
    #     doc.add_heading(f'2.3 托管listing明细（{date2_date}-{date1_date}）：', 3)
    #     csv_path5 = DbToolsCsv(brand, market).get_managed_listing_ditial_data(market, start_date, end_date, asin_info)
    #     insert_csv_to_docx(doc, csv_path5, 8)
    #     json3 = csv_to_json(csv_path5)
    #     translate_kw5 = create_summery(json3, 18)
    #     translate_kw6 = create_summery(json3, 7)
    #     doc.add_paragraph().add_run(f"1.{translate_kw5}")
    #     doc.add_paragraph().add_run(f"2.{translate_kw6}")
    #     doc.add_heading(f'2.4 库存情况（{date2_date}-{date1_date}）：', 3)
    #     csv_path6 = DbToolsCsv(brand, market).get_inventory_data_custody(market, start_date, end_date, asin_info)
    #     insert_csv_to_docx(doc, csv_path6, 8)


    doc.add_heading('3. 总结', 2)
    # result1 = translate_kw1 + "\n" + translate_kw2
    # translate_kw4 = create_summery(result1, 3)
    doc.add_heading('3.1 本月总结（整体总结）', 3)
    doc.add_heading('3.2 下月优化策略', 3)
    doc.add_heading('3.3 下月计划与目标', 3)
    doc.add_heading('3.3.1 下月目标设定', 4)
    doc.add_heading('3.3.2 行动计划', 4)
    doc.add_heading('4. 异常情况及客户协助', 2)


    # sections = doc.sections
    # print(sections)
    # default_section = sections[0]
    # print(default_section.top_margin.cm)  # 2.54
    # print(default_section.right_margin.cm)  # 3.175
    # print(default_section.bottom_margin.cm)  # 2.54
    # print(default_section.left_margin.cm)  # 3.175
    # default_section.top_margin = Cm(2.5)
    # default_section.right_margin = Cm(2)
    # default_section.bottom_margin = Cm(2.5)
    # default_section.left_margin = Cm(2)
    # # 增加分页符
    # doc.add_page_break()
    #
    # # 增加标题 API 分析， 只能设置 0-9 级标题
    # for i in range(0,10):
    #     doc.add_heading(f'标题{i}', i)

    docx_path = f'{brand}_{market}_2024年{week_number}月_({date2_date}-{date1_date}).docx'
    docx_path = os.path.join(get_output_path(), docx_path)
    # 保存文件
    doc.save(docx_path)
    return docx_path

if __name__ == "__main__":
    generate_docx('amazon_youniverse_eu', 'eu','DE', '2024-08-09', '2024-08-28', 20)
    # generate_summary_docx('DELOMO', 'EU', '2024-08-23', '2024-09-22', 31)
