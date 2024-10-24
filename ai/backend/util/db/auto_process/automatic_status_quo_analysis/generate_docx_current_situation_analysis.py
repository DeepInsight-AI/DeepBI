# 导入模块
import asyncio
import json
import os
from datetime import datetime, timedelta
from docx2pdf import convert
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.util.db_tool import DbToolsCsv
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.util.util import csv_to_json
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.util.agent import ask_question
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.output.output_path import get_output_path
from ai.backend.util.db.auto_process.automatic_status_quo_analysis.util.change import convert_to_pdf


def set_font_size(cell, pt_size):
    """设置单元格字体大小"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(pt_size)


def get_font_size1(value):
    """根据float值的长度确定字体大小"""
    if isinstance(value, float):
        length = len(str(int(value)))
        if length >= 10:
            return 6
        elif length == 9:
            return 7
        elif length == 8:
            return 8
        elif length <= 7:
            return 9
        else:
            return 6
    return 6


def get_font_size2(value):
    """根据float值的长度确定字体大小"""
    if isinstance(value, float):
        length = len(str(int(value)))
        if length >= 9:
            return 6
        elif length == 8:
            return 7
        elif length == 7:
            return 7
        elif length <= 6:
            return 8
        else:
            return 6
    return 6


def determine_max_font_size(csv_path,code):
    """读取CSV文件并返回根据float数据长度确定的最大字体大小"""
    df = pd.read_csv(csv_path, encoding='utf-8-sig')

    max_font_size = 9  # 默认字体大小

    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            value = df.iloc[i, j]
            if isinstance(value, (float, np.float64)) and not pd.isna(value):
                if code == 1:
                    font_size = get_font_size1(value)
                elif code == 2:
                    font_size = get_font_size2(value)
                if font_size < max_font_size:
                    max_font_size = font_size
    print(max_font_size)
    return max_font_size


def insert_csv_to_docx(doc, csv_path, pt):
    """将CSV文件插入到给定的docx文档中"""
    df = pd.read_csv(csv_path, encoding='utf-8-sig')
    row = df.shape[0] + 1  # 行数，加标题栏
    col = df.shape[1]  # 列数

    # 设置表格样式
    table = doc.add_table(rows=row, cols=col, style="Table Grid")

    # 写入列标签
    for i in range(col):
        cell = table.cell(0, i)
        cell.text = str(df.columns[i])
        set_font_size(cell, pt)

        # 写入数据
    for i in range(1, row):
        for j in range(col):
            cell = table.cell(i, j)
            value = df.iloc[i - 1, j]
            if pd.isna(value):  # 检查是否为 NaN
                cell.text = '0%'
            elif isinstance(value, str):
                cell.text = value
            elif isinstance(value, (int, np.integer)):
                cell.text = str(int(value))  # 确保整数类型为 int
            elif isinstance(value, float):
                cell.text = '{:.0f}'.format(value)
            set_font_size(cell, pt)
    # #     # 设置列宽以适应内容
    # # for col_idx in range(col):
    # #     max_length = max(len(table.cell(row_idx, col_idx).text) for row_idx in range(row))
    # #     width = Pt(max_length * 12)  # 根据需要调整乘数以增加或减少宽度
    # #     for row_idx in range(row):
    # #         table.cell(row_idx, col_idx).width = width
    # #
    # # 设置字体大小
    # for cell in table._cells:
    #     cell.paragraphs[0].runs[0].font.size = Pt(pt)
    table.autofit = True  # 表格自动适应窗口大小
    table.rows.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    # table.style.font.name = u'楷体'  # 设置字体格式
    # table.style.font.size = Pt(pt)  # 设置字体大小
    return table

def create_summery(date,code):
    summery = asyncio.get_event_loop().run_until_complete(ask_question(date, code))
    return summery
    # return 'test'


def generate_docx(db, brand, market,start_date,end_date):
    # 新建文档对象按模板新建 word 文档文件，具有模板文件的所有格式
    doc = Document()

    # end_date = (datetime.today() - timedelta(days=2)).strftime('%m.%d')
    # start_date = (datetime.today() - timedelta(days=31)).strftime('%m.%d')
    # 增加标题：add_heading(self, text="", level=1):
    date1 = datetime.strptime(end_date, '%Y-%m-%d')
    date2 = datetime.strptime(start_date, '%Y-%m-%d')
    difference = (date1 - date2).days+1
    doc.add_heading(f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析',0)

    doc.add_heading('一、宏观分析', 1)
    doc.add_paragraph().add_run(f'近{difference}天的广告数据和店铺营收数据如下所示:')
    csv_path1 = DbToolsCsv(db, brand, market).get_advertising_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path1,9+1)
    doc.add_paragraph().add_run('')
    csv_path2 = DbToolsCsv(db, brand, market).get_store_data(market,end_date,start_date)
    table1 = insert_csv_to_docx(doc, csv_path2,9+1)
    table1.cell(1, 0).text = ""
    json1 = csv_to_json(csv_path1)
    json2 = csv_to_json(csv_path2)
    translate_kw1 = create_summery(json1, 0)
    translate_kw2 = create_summery(json2, 1)
    translate_kw3 = create_summery(json2, 2)
    result1 = translate_kw1 + "\n" + translate_kw2 + "\n" + translate_kw3
    translate_kw4 = create_summery(result1, 3)
    # doc.add_paragraph().add_run('')
    doc.add_heading('1.1 ACOS分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw1}")
    doc.add_heading('1.2 自然销量占比分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw2}")
    doc.add_heading('1.3 广告花费占比分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw3}")

    doc.add_heading(f'二、现状分析（{start_date}-{end_date}，{difference}天）', 1)
    doc.add_heading(f'1.目前所有在投计划情况（{start_date}-{end_date}）', 2)
    csv_path3 = DbToolsCsv(db, brand, market).get_ad_type(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path3,9+1)
    doc.add_paragraph().add_run('')
    csv_path4 = DbToolsCsv(db, brand, market).get_ad_type_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path4,9+1)
    doc.add_paragraph().add_run('')
    csv_path5 = DbToolsCsv(db, brand, market).get_sp_type_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path5,9+1)
    # doc.add_paragraph().add_run('')
    json4 = csv_to_json(csv_path4)
    json5 = csv_to_json(csv_path5)
    translate_kw5 = create_summery(json4, 4)
    translate_kw6 = create_summery(json5, 5)
    doc.add_heading('2.1.1 整体销售额占比分析', 3)
    doc.add_paragraph().add_run(f"{translate_kw5}")
    doc.add_heading('2.1.2 SP广告销售额占比分析', 3)
    doc.add_paragraph().add_run(f"{translate_kw6}")
    # doc.add_paragraph().add_run('')
    doc.add_heading(f'2.以下是按父Asin分类，各listing的数据情况', 2)
    csv_path6 = DbToolsCsv(db, brand, market).get_listing_summary_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path6,8+1+1)
    doc.add_paragraph().add_run('')
    doc.add_heading(f'2.2.1 SP计划整体', 3)
    csv_path7 = DbToolsCsv(db, brand, market).get_listing_sp_summary_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path7,8+1+1)
    doc.add_paragraph().add_run('')
    json7 = csv_to_json(csv_path7)
    translate_kw7 = create_summery(json7, 6)
    translate_kw8 = create_summery(json7, 7)
    doc.add_heading('2.2.1.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw7}")
    doc.add_heading('2.2.1.2 整体销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw8}")
    doc.add_heading(f'2.2.2 SP手动和SP自动计划', 3)
    csv_path8 = DbToolsCsv(db, brand, market).get_listing_sp_specific_data(market,end_date,start_date)
    pt1 = determine_max_font_size(csv_path8,1)
    insert_csv_to_docx(doc, csv_path8,pt1)
    # doc.add_paragraph().add_run('')
    json8 = csv_to_json(csv_path8)
    translate_kw9 = create_summery(json8, 8)
    translate_kw10 = create_summery(json8, 9)
    doc.add_heading('2.2.2.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw9}")
    doc.add_heading('2.2.2.2 SP广告销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw10}")
    doc.add_heading(f'2.2.3 SD计划', 3)
    csv_path9 = DbToolsCsv(db, brand, market).get_listing_sd_summary_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path9,9+1)
    json9 = csv_to_json(csv_path9)
    translate_kw11 = create_summery(json9, 10)
    translate_kw12 = create_summery(json9, 11)
    doc.add_heading('2.2.3.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw12}")
    doc.add_heading('2.2.3.2 SD广告销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw11}")

    doc.add_heading(f'三、目标期望设定', 1)
    doc.add_paragraph().add_run('按照我们先前的经验，广告营收大致占总营收的45-55%；而在广告营收中，SD广告和SP广告带来的营收占比分别为25%-35%和65%-75%。于是我们的长期目标是希望广告整体营收占比可以达到50%，整体Acos值可以降低到24%以内，广告花费占比降低至12%。')
    doc.add_paragraph().add_run('为没有开设SD广告的商品开设SD广告，并将SD广告营收占比提升至30%，Acos值控制在8%以内；SP广告营收占比70%，Acos值控制在24%以内。')
    doc.add_paragraph().add_run('于是，我们将每个listing的SD、SP期望广告营收占比分别设置为30%和70%，期望Acos值设置为8%和24%。（若目前已达到预期目标，则将目前的值设置为预期目标）')
    doc.add_heading(f'广告销售额期望计算', 2)
    doc.add_paragraph().add_run('对于每个listing，我们将以SP或SD中较好的作为基准，按照目标比例提升另一类型的广告数据。结算结果如下所示')
    csv_path10,expect_ad_sales,total_ad_sales = DbToolsCsv(db, brand, market).get_expected_sales(market,end_date,start_date)
    pt2 = determine_max_font_size(csv_path10, 2)
    insert_csv_to_docx(doc, csv_path10,pt2)
    doc.add_paragraph().add_run('')
    doc.add_paragraph().add_run(f'广告销售额上可达到（{expect_ad_sales}/{total_ad_sales}-1）*100%={round((expect_ad_sales/total_ad_sales-1)*100,2)}%的增幅。')
    # doc.add_heading(f'2、广告成本期望计算', 2)
    csv_path11, expect_ad_cost, total_ad_cost = DbToolsCsv(db, brand, market).get_expected_cost(market,end_date,start_date)
    # insert_csv_to_docx(doc, csv_path11,7)
    # doc.add_paragraph().add_run(f'整体广告花费下降（1-{expect_ad_cost}/{total_ad_cost}）*100%={round((1 - expect_ad_cost / total_ad_cost) * 100, 2)}%')
    doc.add_paragraph().add_run(f'整体预期AOCS值为（{expect_ad_cost}/{expect_ad_sales}）*100%={round((expect_ad_cost / expect_ad_sales) * 100, 2)}%')
    doc.add_paragraph().add_run(f'我们预期自然销售占比在45%-55%之间，根据这个比例，则预期Tacos值为{round((expect_ad_cost / expect_ad_sales) * 45, 2)}%-{round((expect_ad_cost / expect_ad_sales) * 55, 2)}%')

    doc.add_heading(f'四、总结', 1)
    doc.add_paragraph().add_run(f"1.{translate_kw4}")
    result2 = translate_kw11 + "\n" + translate_kw12
    result3 = translate_kw9 + "\n" + translate_kw10
    translate_kw13 = create_summery(result2, 12)
    translate_kw14 = create_summery(result3, 13)
    doc.add_paragraph().add_run(f"2.{translate_kw13}")
    doc.add_paragraph().add_run(f"3.{translate_kw14}")
    doc.add_paragraph().add_run(f'4.最终实现，预期销售额{round((expect_ad_sales/total_ad_sales-1)*100,2)}%的增长，预期整体ACOS为{round((expect_ad_cost / expect_ad_sales) * 100, 2)}%，预期Tacos值为{round((expect_ad_cost / expect_ad_sales) * 45, 2)}%-{round((expect_ad_cost / expect_ad_sales) * 55, 2)}%的目标')
    sections = doc.sections
    print(sections)
    default_section = sections[0]
    print(default_section.top_margin.cm)  # 2.54
    print(default_section.right_margin.cm)  # 3.175
    print(default_section.bottom_margin.cm)  # 2.54
    print(default_section.left_margin.cm)  # 3.175
    default_section.top_margin = Cm(2.5)
    default_section.right_margin = Cm(2)
    default_section.bottom_margin = Cm(2.5)
    default_section.left_margin = Cm(2)
    # # 增加分页符
    # doc.add_page_break()
    #
    # # 增加标题 API 分析， 只能设置 0-9 级标题
    # for i in range(0,10):
    #     doc.add_heading(f'标题{i}', i)
    # for paragraph in doc.paragraphs:
    #     for run in paragraph.runs:
    #         run.font.name = 'Arial'  # 设置字体
    #         run.font.size = Pt(12)  # 设置字号

    docx_path = f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析.docx'
    pdf_path = f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析.pdf'
    # 保存文件
    docx_path = os.path.join(get_output_path(), docx_path)

    # Save DOCX file
    doc.save(docx_path)

    # Convert DOCX to PDF
    convert(docx_path)
    pdf_path = os.path.join(get_output_path(), pdf_path)
    # convert_to_pdf(docx_path,pdf_path)
    return pdf_path


def generate_docx_supplier(db, brand, market,start_date,end_date):
    # 新建文档对象按模板新建 word 文档文件，具有模板文件的所有格式
    doc = Document()
    # end_date = (datetime.today() - timedelta(days=2)).strftime('%m.%d')
    # start_date = (datetime.today() - timedelta(days=31)).strftime('%m.%d')
    # 增加标题：add_heading(self, text="", level=1):
    date1 = datetime.strptime(end_date, '%Y-%m-%d')
    date2 = datetime.strptime(start_date, '%Y-%m-%d')
    difference = (date1 - date2).days+1
    doc.add_heading(f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析',0)

    doc.add_heading('一、宏观分析', 1)
    doc.add_paragraph().add_run(f'近{difference}天的广告数据和店铺营收数据如下所示:')
    csv_path1 = DbToolsCsv(db, brand, market).get_advertising_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path1,9+1)
    doc.add_paragraph().add_run('')
    csv_path2 = DbToolsCsv(db, brand, market).get_store_data_supplier(market,end_date,start_date)
    table1 = insert_csv_to_docx(doc, csv_path2,9+1)
    table1.cell(1, 0).text = ""
    json1 = csv_to_json(csv_path1)
    json2 = csv_to_json(csv_path2)
    translate_kw1 = create_summery(json1, 0)
    translate_kw2 = create_summery(json2, 1)
    translate_kw3 = create_summery(json2, 2)
    result1 = translate_kw1 + "\n" + translate_kw2 + "\n" + translate_kw3
    translate_kw4 = create_summery(result1, 3)
    # doc.add_paragraph().add_run('')
    doc.add_heading('1.1 ACOS分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw1}")
    doc.add_heading('1.2 自然销量占比分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw2}")
    doc.add_heading('1.3 广告花费占比分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw3}")

    doc.add_heading(f'二、现状分析（{start_date}-{end_date}，{difference}天）', 1)
    doc.add_heading(f'1.目前所有在投计划情况（{start_date}-{end_date}）', 2)
    csv_path3 = DbToolsCsv(db, brand, market).get_ad_type(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path3,9+1)
    doc.add_paragraph().add_run('')
    csv_path4 = DbToolsCsv(db, brand, market).get_ad_type_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path4,9+1)
    doc.add_paragraph().add_run('')
    csv_path5 = DbToolsCsv(db, brand, market).get_sp_type_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path5,9+1)
    # doc.add_paragraph().add_run('')
    json4 = csv_to_json(csv_path4)
    json5 = csv_to_json(csv_path5)
    translate_kw5 = create_summery(json4, 4)
    translate_kw6 = create_summery(json5, 5)
    doc.add_heading('2.1.1 整体销售额占比分析', 3)
    doc.add_paragraph().add_run(f"{translate_kw5}")
    doc.add_heading('2.1.2 SP广告销售额占比分析', 3)
    doc.add_paragraph().add_run(f"{translate_kw6}")
    # doc.add_paragraph().add_run('')
    doc.add_heading(f'2.以下是按父Asin分类，各listing的数据情况', 2)
    csv_path6 = DbToolsCsv(db, brand, market).get_listing_summary_data_supplier(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path6,8+1+1)
    doc.add_paragraph().add_run('')
    doc.add_heading(f'2.2.1 SP计划整体', 3)
    csv_path7 = DbToolsCsv(db, brand, market).get_listing_sp_summary_data_supplier(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path7,8+1+1)
    doc.add_paragraph().add_run('')
    json7 = csv_to_json(csv_path7)
    translate_kw7 = create_summery(json7, 6)
    translate_kw8 = create_summery(json7, 7)
    doc.add_heading('2.2.1.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw7}")
    doc.add_heading('2.2.1.2 整体销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw8}")
    doc.add_heading(f'2.2.2 SP手动和SP自动计划', 3)
    csv_path8 = DbToolsCsv(db, brand, market).get_listing_sp_specific_data_supplier(market,end_date,start_date)
    pt1 = determine_max_font_size(csv_path8,1)
    insert_csv_to_docx(doc, csv_path8,pt1)
    # doc.add_paragraph().add_run('')
    json8 = csv_to_json(csv_path8)
    translate_kw9 = create_summery(json8, 8)
    translate_kw10 = create_summery(json8, 9)
    doc.add_heading('2.2.2.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw9}")
    doc.add_heading('2.2.2.2 SP广告销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw10}")
    doc.add_heading(f'2.2.3 SD计划', 3)
    csv_path9 = DbToolsCsv(db, brand, market).get_listing_sd_summary_data_supplier(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path9,9+1)
    json9 = csv_to_json(csv_path9)
    translate_kw11 = create_summery(json9, 10)
    translate_kw12 = create_summery(json9, 11)
    doc.add_heading('2.2.3.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw12}")
    doc.add_heading('2.2.3.2 SD广告销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw11}")

    doc.add_heading(f'三、目标期望设定', 1)
    doc.add_paragraph().add_run('按照我们先前的经验，广告营收大致占总营收的50-60%；而在广告营收中，SD广告和SP广告带来的营收占比分别为35%-40%和50%-65%。于是我们的长期目标是希望广告整体营收占比可以达到55%，整体Acos值可以降低到24%以内，广告花费占比降低至12%。')
    doc.add_paragraph().add_run('为没有开设SD广告的商品开设SD广告，并将SD广告营收占比提升至35%，Acos值控制在8%以内；SP广告营收占比65%，Acos值控制在24%以内。')
    doc.add_paragraph().add_run('于是，我们将每个listing的SD、SP期望广告营收占比分别设置为35%和65%，期望Acos值设置为8%和24%。（若目前已达到预期目标，则将目前的值设置为预期目标）')
    doc.add_heading(f'广告销售额期望计算', 2)
    doc.add_paragraph().add_run('对于每个listing，我们将以SP或SD中较好的作为基准，按照目标比例提升另一类型的广告数据。结算结果如下所示')
    csv_path10,expect_ad_sales,total_ad_sales = DbToolsCsv(db, brand, market).get_expected_sales_supplier(market,end_date,start_date)
    pt2 = determine_max_font_size(csv_path10, 2)
    insert_csv_to_docx(doc, csv_path10,pt2)
    doc.add_paragraph().add_run('')
    doc.add_paragraph().add_run(f'广告销售额上可达到（{expect_ad_sales}/{total_ad_sales}-1）*100%={round((expect_ad_sales/total_ad_sales-1)*100,2)}%的增幅。')
    # doc.add_heading(f'2、广告成本期望计算', 2)
    csv_path11, expect_ad_cost, total_ad_cost = DbToolsCsv(db, brand, market).get_expected_cost_supplier(market,end_date,start_date)
    # insert_csv_to_docx(doc, csv_path11,7)
    # doc.add_paragraph().add_run(f'整体广告花费下降（1-{expect_ad_cost}/{total_ad_cost}）*100%={round((1 - expect_ad_cost / total_ad_cost) * 100, 2)}%')
    doc.add_paragraph().add_run(f'整体预期AOCS值为（{expect_ad_cost}/{expect_ad_sales}）*100%={round((expect_ad_cost / expect_ad_sales) * 100, 2)}%')
    doc.add_paragraph().add_run(f'我们预期自然销售占比在40%-50%之间，根据这个比例，则预期Tacos值为{round((expect_ad_cost / expect_ad_sales) * 50, 2)}%-{round((expect_ad_cost / expect_ad_sales) * 60, 2)}%')

    doc.add_heading(f'四、总结', 1)
    doc.add_paragraph().add_run(f"1.{translate_kw4}")
    result2 = translate_kw11 + "\n" + translate_kw12
    result3 = translate_kw9 + "\n" + translate_kw10
    translate_kw13 = create_summery(result2, 12)
    translate_kw14 = create_summery(result3, 13)
    doc.add_paragraph().add_run(f"2.{translate_kw13}")
    doc.add_paragraph().add_run(f"3.{translate_kw14}")
    doc.add_paragraph().add_run(f'4.最终实现，预期销售额{round((expect_ad_sales/total_ad_sales-1)*100,2)}%的增长，预期整体ACOS为{round((expect_ad_cost / expect_ad_sales) * 100, 2)}%，预期Tacos值为{round((expect_ad_cost / expect_ad_sales) * 50, 2)}%-{round((expect_ad_cost / expect_ad_sales) * 60, 2)}%的目标')
    sections = doc.sections
    print(sections)
    default_section = sections[0]
    print(default_section.top_margin.cm)  # 2.54
    print(default_section.right_margin.cm)  # 3.175
    print(default_section.bottom_margin.cm)  # 2.54
    print(default_section.left_margin.cm)  # 3.175
    default_section.top_margin = Cm(2.5)
    default_section.right_margin = Cm(2)
    default_section.bottom_margin = Cm(2.5)
    default_section.left_margin = Cm(2)
    # # 增加分页符
    # doc.add_page_break()
    #
    # # 增加标题 API 分析， 只能设置 0-9 级标题
    # for i in range(0,10):
    #     doc.add_heading(f'标题{i}', i)

    docx_path = f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析.docx'
    pdf_path = f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析.pdf'
    # 保存文件
    docx_path = os.path.join(get_output_path(), docx_path)

    # Save DOCX file
    doc.save(docx_path)

    # Convert DOCX to PDF
    convert(docx_path)
    pdf_path = os.path.join(get_output_path(), pdf_path)
    return pdf_path


def generate_docx_test(db,brand, market,start_date,end_date):
    # 新建文档对象按模板新建 word 文档文件，具有模板文件的所有格式
    doc = Document()
    # end_date = (datetime.today() - timedelta(days=2)).strftime('%m.%d')
    # start_date = (datetime.today() - timedelta(days=31)).strftime('%m.%d')
    # 增加标题：add_heading(self, text="", level=1):
    date1 = datetime.strptime(end_date, '%Y-%m-%d')
    date2 = datetime.strptime(start_date, '%Y-%m-%d')
    difference = (date1 - date2).days+1
    doc.add_heading(f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析',0)

    doc.add_heading('一、宏观分析', 1)
    doc.add_paragraph().add_run(f'近{difference}天的广告数据和店铺营收数据如下所示:')
    csv_path1 = DbToolsCsv(db,brand, market).get_advertising_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path1,9+1)
    doc.add_paragraph().add_run('')
    csv_path2 = DbToolsCsv(db,brand, market).get_store_data(market,end_date,start_date)
    table1 = insert_csv_to_docx(doc, csv_path2,9+1)
    #table1.cell(1, 0).text = ""
    json1 = csv_to_json(csv_path1)
    json2 = csv_to_json(csv_path2)
    translate_kw1 = create_summery(json1, 0)
    translate_kw2 = create_summery(json2, 1)
    translate_kw3 = create_summery(json2, 2)
    result1 = translate_kw1 + "\n" + translate_kw2 + "\n" + translate_kw3
    translate_kw4 = create_summery(result1, 3)
    # doc.add_paragraph().add_run('')
    doc.add_heading('1.1 ACOS分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw1}")
    doc.add_heading('1.2 自然销量占比分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw2}")
    doc.add_heading('1.3 广告花费占比分析', 2)
    doc.add_paragraph().add_run(f"{translate_kw3}")

    doc.add_heading(f'二、现状分析（{start_date}-{end_date}，{difference}天）', 1)
    doc.add_heading(f'1.目前所有在投计划情况（{start_date}-{end_date}）', 2)
    csv_path3 = DbToolsCsv(db,brand, market).get_ad_type(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path3,9+1)
    doc.add_paragraph().add_run('')
    csv_path4 = DbToolsCsv(db,brand, market).get_ad_type_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path4,9+1)
    doc.add_paragraph().add_run('')
    csv_path5 = DbToolsCsv(db,brand, market).get_sp_type_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path5,9+1)
    # doc.add_paragraph().add_run('')
    json4 = csv_to_json(csv_path4)
    json5 = csv_to_json(csv_path5)
    translate_kw5 = create_summery(json4, 4)
    translate_kw6 = create_summery(json5, 5)
    doc.add_heading('2.1.1 整体销售额占比分析', 3)
    doc.add_paragraph().add_run(f"{translate_kw5}")
    doc.add_heading('2.1.2 SP广告销售额占比分析', 3)
    doc.add_paragraph().add_run(f"{translate_kw6}")
    # doc.add_paragraph().add_run('')
    doc.add_heading(f'2.以下是按父Asin分类，各listing的数据情况', 2)
    csv_path6 = DbToolsCsv(db,brand, market).get_listing_summary_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path6,8+1+1)
    doc.add_paragraph().add_run('')
    doc.add_heading(f'2.2.1 SP计划整体', 3)
    csv_path7 = DbToolsCsv(db,brand, market).get_listing_sp_summary_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path7,8+1+1)
    doc.add_paragraph().add_run('')
    json7 = csv_to_json(csv_path7)
    translate_kw7 = create_summery(json7, 6)
    translate_kw8 = create_summery(json7, 7)
    doc.add_heading('2.2.1.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw7}")
    doc.add_heading('2.2.1.2 整体销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw8}")
    doc.add_heading(f'2.2.2 SP手动和SP自动计划', 3)
    csv_path8 = DbToolsCsv(db,brand, market).get_listing_sp_specific_data(market,end_date,start_date)
    pt1 = determine_max_font_size(csv_path8,1)
    insert_csv_to_docx(doc, csv_path8,pt1)
    # doc.add_paragraph().add_run('')
    json8 = csv_to_json(csv_path8)
    translate_kw9 = create_summery(json8, 8)
    translate_kw10 = create_summery(json8, 9)
    doc.add_heading('2.2.2.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw9}")
    doc.add_heading('2.2.2.2 SP广告销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw10}")
    doc.add_heading(f'2.2.3 SD计划', 3)
    csv_path9 = DbToolsCsv(db,brand, market).get_listing_sd_summary_data(market,end_date,start_date)
    insert_csv_to_docx(doc, csv_path9,9+1)
    json9 = csv_to_json(csv_path9)
    translate_kw11 = create_summery(json9, 10)
    translate_kw12 = create_summery(json9, 11)
    doc.add_heading('2.2.3.1 ACOS分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw12}")
    doc.add_heading('2.2.3.2 SD广告销售额占比分析', 4)
    doc.add_paragraph().add_run(f"{translate_kw11}")

    doc.add_heading(f'三、目标期望设定', 1)
    doc.add_paragraph().add_run('按照我们先前的经验，广告营收大致占总营收的55-60%；而在广告营收中，SD广告和SP广告带来的营收占比分别为35%-40%和50%-65%。于是我们的长期目标是希望广告整体营收占比可以达到55%，整体Acos值可以降低到24%以内，广告花费占比降低至12%。')
    doc.add_paragraph().add_run('为没有开设SD广告的商品开设SD广告，并将SD广告营收占比提升至35%，Acos值控制在8%以内；SP广告营收占比65%，Acos值控制在24%以内。')
    doc.add_paragraph().add_run('于是，我们将每个listing的SD、SP期望广告营收占比分别设置为35%和65%，期望Acos值设置为8%和24%。（若目前已达到预期目标，则将目前的值设置为预期目标）')
    doc.add_heading(f'广告销售额期望计算', 2)
    doc.add_paragraph().add_run('对于每个listing，我们将以SP或SD中较好的作为基准，按照目标比例提升另一类型的广告数据。结算结果如下所示')
    csv_path10,expect_ad_sales,total_ad_sales = DbToolsCsv(db,brand, market).get_expected_sales(market,end_date,start_date)
    pt2 = determine_max_font_size(csv_path10, 2)
    insert_csv_to_docx(doc, csv_path10,pt2)
    doc.add_paragraph().add_run('')
    # doc.add_paragraph().add_run(f'广告销售额上可达到（{expect_ad_sales}/{total_ad_sales}-1）*100%={round((expect_ad_sales/total_ad_sales-1)*100,2)}%的增幅。')
    # doc.add_heading(f'2、广告成本期望计算', 2)
    csv_path11, expect_ad_cost, total_ad_cost = DbToolsCsv(db,brand, market).get_expected_cost(market,end_date,start_date)
    # insert_csv_to_docx(doc, csv_path11,7)
    # doc.add_paragraph().add_run(f'整体广告花费下降（1-{expect_ad_cost}/{total_ad_cost}）*100%={round((1 - expect_ad_cost / total_ad_cost) * 100, 2)}%')
    # doc.add_paragraph().add_run(f'整体预期AOCS值为（{expect_ad_cost}/{expect_ad_sales}）*100%={round((expect_ad_cost / expect_ad_sales) * 100, 2)}%')
    # doc.add_paragraph().add_run(f'我们预期自然销售占比在40%-50%之间，根据这个比例，则预期Tacos值为{round((expect_ad_cost / expect_ad_sales) * 50, 2)}%-{round((expect_ad_cost / expect_ad_sales) * 60, 2)}%')

    doc.add_heading(f'四、总结', 1)
    doc.add_paragraph().add_run(f"1.{translate_kw4}")
    result2 = translate_kw11 + "\n" + translate_kw12
    result3 = translate_kw9 + "\n" + translate_kw10
    translate_kw13 = create_summery(result2, 12)
    translate_kw14 = create_summery(result3, 13)
    doc.add_paragraph().add_run(f"2.{translate_kw13}")
    doc.add_paragraph().add_run(f"3.{translate_kw14}")
    # doc.add_paragraph().add_run(f'4.最终实现，预期销售额{round((expect_ad_sales/total_ad_sales-1)*100,2)}%的增长，预期整体ACOS为{round((expect_ad_cost / expect_ad_sales) * 100, 2)}%，预期Tacos值为{round((expect_ad_cost / expect_ad_sales) * 50, 2)}%-{round((expect_ad_cost / expect_ad_sales) * 60, 2)}%的目标')
    sections = doc.sections
    print(sections)
    default_section = sections[0]
    print(default_section.top_margin.cm)  # 2.54
    print(default_section.right_margin.cm)  # 3.175
    print(default_section.bottom_margin.cm)  # 2.54
    print(default_section.left_margin.cm)  # 3.175
    default_section.top_margin = Cm(2.5)
    default_section.right_margin = Cm(2)
    default_section.bottom_margin = Cm(2.5)
    default_section.left_margin = Cm(2)
    # # 增加分页符
    # doc.add_page_break()
    #
    # # 增加标题 API 分析， 只能设置 0-9 级标题
    # for i in range(0,10):
    #     doc.add_heading(f'标题{i}', i)

    docx_path = f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析.docx'
    pdf_path = f'{brand}_{market}_{start_date}-{end_date}({difference}天)_现状分析.pdf'
    # 保存文件
    docx_path = os.path.join(get_output_path(), docx_path)

    # Save DOCX file
    doc.save(docx_path)

    # Convert DOCX to PDF
    convert(docx_path)
    pdf_path = os.path.join(get_output_path(), pdf_path)
    return pdf_path

if __name__ == "__main__":
    generate_docx('amazon_kfeiya','COFaR','US','2024-09-11','2024-10-10')
    # generate_docx_supplier('LAPASA','US','2024-08-10','2024-09-08')
